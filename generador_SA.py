import streamlit as st
from docx import Document
from fpdf import FPDF
import os

def generar_word(datos):
    doc = Document()
    doc.add_heading('Situación de Aprendizaje', level=1)
    
    for clave, valor in datos.items():
        doc.add_heading(clave, level=2)
        doc.add_paragraph(valor)
    
    nombre_archivo = "situacion_aprendizaje.docx"
    doc.save(nombre_archivo)
    return nombre_archivo

def generar_pdf(datos):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", style='B', size=16)
    pdf.cell(200, 10, 'Situación de Aprendizaje', ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_font("Arial", size=12)
    for clave, valor in datos.items():
        pdf.set_font("Arial", style='B', size=12)
        pdf.cell(200, 10, clave, ln=True)
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, valor)
        pdf.ln(5)
    
    nombre_archivo = "situacion_aprendizaje.pdf"
    pdf.output(nombre_archivo)
    return nombre_archivo

def main():
    st.title("Generador de Situaciones de Aprendizaje")
    st.subheader("Formulario de entrada")
    
    datos = {}
    datos["Etapa educativa"] = st.text_input("Etapa educativa")
    datos["Área o materia"] = st.text_input("Área o materia")
    datos["Objetivos de aprendizaje"] = st.text_area("Objetivos de aprendizaje")
    datos["Competencias clave"] = st.text_area("Competencias clave")
    datos["Criterios de evaluación"] = st.text_area("Criterios de evaluación")
    datos["Metodología y actividades"] = st.text_area("Metodología y actividades")
    datos["Evaluación y rúbricas"] = st.text_area("Evaluación y rúbricas")
    
    if st.button("Generar documentos"):
        word_file = generar_word(datos)
        pdf_file = generar_pdf(datos)
        
        with open(word_file, "rb") as f:
            st.download_button("Descargar Word", f, file_name=word_file)
        with open(pdf_file, "rb") as f:
            st.download_button("Descargar PDF", f, file_name=pdf_file)

if __name__ == "__main__":
    main()
